from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tax_doc():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Viva Guide: How the Indian Tax System Works', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Use this document to explain the logic behind your project to the external examiner.")

    # --- Section 1 ---
    doc.add_heading('1. The Two Regimes', level=1)
    p = doc.add_paragraph()
    p.add_run('Currently, India has two tax systems. A taxpayer can choose either:\n').bold = True
    doc.add_paragraph('• The New Regime (Default): Lower tax rates, but NO deductions (no 80C, HRA, etc.). This is the logic used in this project.')
    doc.add_paragraph('• The Old Regime: Higher tax rates, but allows deductions for investments.')

    # --- Section 2 ---
    doc.add_heading('2. The 5 Heads of Income', level=1)
    doc.add_paragraph('In reality, income is classified into 5 categories. Your project simplifies this into "Salary" and "Other Sources":')
    items = [
        'Salary (Employment income)',
        'House Property (Rent received)',
        'Profits from Business/Profession',
        'Capital Gains (Stock market/Real estate profits)',
        'Other Sources (Bank interest, lottery, etc.)'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # --- Section 3 ---
    doc.add_heading('3. The Calculation Algorithm (New Regime)', level=1)
    doc.add_paragraph("This is the mathematical logic used in your 'utils.py' file:")
    
    step1 = doc.add_paragraph()
    step1.add_run('Step 1: Gross Total Income').bold = True
    step1.add_run(' = Salary + Other Income')
    
    step2 = doc.add_paragraph()
    step2.add_run('Step 2: Deductions').bold = True
    step2.add_run(' = Subtract Standard Deduction (₹50,000 or ₹75,000 depending on budget year).')
    
    step3 = doc.add_paragraph()
    step3.add_run('Step 3: Net Taxable Income').bold = True
    step3.add_run(' = Gross Income - Deductions')

    step4 = doc.add_paragraph()
    step4.add_run('Step 4: Apply Slabs').bold = True
    doc.add_paragraph('India uses a Progressive Slab System. You pay different rates for different chunks of money:')

    # --- Table ---
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Income Slab'
    hdr_cells[1].text = 'Tax Rate (New Regime)'

    data = [
        ('₹0 - ₹3 Lakh', 'Nil (0%)'),
        ('₹3 Lakh - ₹7 Lakh', '5%'),
        ('₹7 Lakh - ₹10 Lakh', '10%'),
        ('₹10 Lakh - ₹12 Lakh', '15%'),
        ('₹12 Lakh - ₹15 Lakh', '20%'),
        ('Above ₹15 Lakh', '30%')
    ]

    for slab, rate in data:
        row_cells = table.add_row().cells
        row_cells[0].text = slab
        row_cells[1].text = rate

    # --- Section 4 ---
    doc.add_heading('4. Viva Presentation Statement', level=1)
    doc.add_paragraph("When asked about the project logic, say this:")
    
    quote = doc.add_paragraph()
    quote.add_run('"My system models the Indian New Tax Regime. It takes the Gross Income, subtracts the Standard Deduction, applies the Progressive Slab Rates (0%, 5%, 10%), and generates a final tax liability receipt, simulating the real-world ITR Filing Process."').italic = True

    # Save
    doc.save('Indian_Tax_System_Viva_Notes.docx')
    print("Success! 'Indian_Tax_System_Viva_Notes.docx' has been created.")

if __name__ == "__main__":
    create_tax_doc()